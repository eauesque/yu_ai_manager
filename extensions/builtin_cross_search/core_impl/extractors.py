"""Text extraction by file format.

Each extraction function takes a file path and returns a (title, content) tuple.
Returns None if the required library is not installed.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_txt(path: Path) -> tuple[str, str]:
    """Plain text."""
    content = path.read_text(encoding="utf-8", errors="replace")
    return path.stem, content


def extract_pdf(path: Path) -> tuple[str, str] | None:
    """Extract text from PDF using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text as pm_extract
        from pdfminer.pdfdocument import PDFDocument
        from pdfminer.pdfparser import PDFParser
    except ImportError:
        logger.debug("pdfminer.six not installed, skipping PDF: %s", path)
        return None

    try:
        # Get title from metadata
        title = path.stem
        try:
            with open(path, "rb") as f:
                parser = PDFParser(f)
                doc = PDFDocument(parser)
                if doc.info:
                    info = doc.info[0] if doc.info else {}
                    raw_title = info.get("Title", b"")
                    if isinstance(raw_title, bytes):
                        raw_title = raw_title.decode("utf-8", errors="replace")
                    raw_title = str(raw_title).strip()
                    if raw_title:
                        title = raw_title
        except Exception:
            logger.debug("PDF title extraction failed", exc_info=True)

        content = pm_extract(str(path))
        return title, content
    except Exception as exc:
        logger.warning("PDF extraction error: %s: %s", path, exc)
        return None


def extract_docx(path: Path) -> tuple[str, str] | None:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
    except ImportError:
        logger.debug("python-docx not installed, skipping DOCX: %s", path)
        return None

    try:
        doc = docx.Document(str(path))
        title = path.stem
        # Try to get title from core_properties
        if doc.core_properties and doc.core_properties.title:
            title = doc.core_properties.title

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return title, content
    except Exception as exc:
        logger.warning("DOCX extraction error: %s: %s", path, exc)
        return None


def extract_rtf(path: Path) -> tuple[str, str] | None:
    """Extract text from RTF using striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        logger.debug("striprtf not installed, skipping RTF: %s", path)
        return None

    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        content = rtf_to_text(raw)
        return path.stem, content
    except Exception as exc:
        logger.warning("RTF extraction error: %s: %s", path, exc)
        return None


# ── Extension -> extractor function mapping ──────────────────────────
EXTRACTORS = {
    ".txt": extract_txt,
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".rtf": extract_rtf,
}

# All extensions to scan
SUPPORTED_EXTENSIONS = frozenset(EXTRACTORS.keys())


def extract_text(path: Path) -> tuple[str, str] | None:
    """Call the appropriate extraction function based on path extension."""
    ext = path.suffix.lower()
    func = EXTRACTORS.get(ext)
    if not func:
        return None
    return func(path)
